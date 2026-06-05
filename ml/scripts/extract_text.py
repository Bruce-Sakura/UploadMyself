#!/usr/bin/env python3
"""
Extract text from uploaded files: PDF, DOCX, images (OCR), plain text.
Usage: python extract_text.py --input <file> [--output <json_file>]
Stdout: JSON {"text": "...", "pages": N, "method": "pdf|docx|ocr|text"}
"""

import argparse
import json
import sys
import os


def extract_pdf(path: str) -> dict:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        # Fallback: pdfplumber
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            return {"text": "\n\n".join(text_parts), "pages": len(text_parts), "method": "pdf"}
        except ImportError:
            return {"error": "No PDF library installed. Install PyMuPDF or pdfplumber."}

    doc = fitz.open(path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return {"text": "\n\n".join(text_parts), "pages": len(text_parts), "method": "pdf"}


def extract_docx(path: str) -> dict:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed."}

    doc = Document(path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return {"text": "\n\n".join(text_parts), "pages": 1, "method": "docx"}


def extract_image_ocr(path: str) -> dict:
    """Extract text from image using Tesseract OCR or PaddleOCR."""
    # Try PaddleOCR first (better for Chinese)
    try:
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang='ch', show_log=False)
        result = ocr.ocr(path, cls=True)
        text_parts = []
        for line in result:
            if line:
                for word_info in line:
                    text_parts.append(word_info[1][0])
        if text_parts:
            return {"text": "\n".join(text_parts), "pages": 1, "method": "paddleocr"}
    except ImportError:
        pass

    # Fallback: Tesseract
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(path)
        text = pytesseract.image_to_string(img, lang='chi_sim+eng')
        return {"text": text, "pages": 1, "method": "tesseract"}
    except ImportError:
        pass

    # Fallback: EasyOCR
    try:
        import easyocr
        reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)
        result = reader.readtext(path)
        text_parts = [item[1] for item in result]
        return {"text": "\n".join(text_parts), "pages": 1, "method": "easyocr"}
    except ImportError:
        pass

    return {"error": "No OCR library installed. Install pytesseract, easyocr, or paddleocr."}


def extract_text(path: str) -> dict:
    """Extract plain text from file."""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(path, 'r', encoding='latin-1') as f:
            text = f.read()
    return {"text": text, "pages": 1, "method": "text"}


def extract_markdown(path: str) -> dict:
    """Extract text from markdown file."""
    return extract_text(path)  # Same as plain text


def main():
    parser = argparse.ArgumentParser(description="Extract text from files")
    parser.add_argument("--input", required=True, help="Input file path")
    parser.add_argument("--output", help="Output JSON file (default: stdout)")
    parser.add_argument("--method", choices=["auto", "pdf", "docx", "ocr", "text"],
                        default="auto", help="Extraction method")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        result = {"error": f"File not found: {args.input}"}
    else:
        ext = os.path.splitext(args.input)[1].lower()
        method = args.method

        if method == "auto":
            if ext == ".pdf":
                method = "pdf"
            elif ext in (".docx", ".doc"):
                method = "docx"
            elif ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"):
                method = "ocr"
            elif ext in (".md", ".markdown"):
                method = "text"
            else:
                method = "text"

        if method == "pdf":
            result = extract_pdf(args.input)
        elif method == "docx":
            result = extract_docx(args.input)
        elif method == "ocr":
            result = extract_image_ocr(args.input)
        else:
            result = extract_text(args.input)

    output_json = json.dumps(result, ensure_ascii=False)

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(output_json)
    else:
        print(output_json)

    if "error" in result:
        sys.exit(1)


if __name__ == "__main__":
    main()
